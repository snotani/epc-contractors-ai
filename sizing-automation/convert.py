#!/usr/bin/env python3
"""
Convert sizing-automation project files to Markdown.

Supported formats:
  - PDF  (.pdf)        → PyMuPDF text extraction
  - Word (.docx)       → python-docx paragraphs + tables
  - Excel (.xlsx/xlsm) → openpyxl sheets → Markdown tables
  - Email (.eml)       → stdlib email parser
  - Email (.msg)       → extract-msg
  - Text (.txt/.log)   → plain copy

Skipped: .dwg, .zip, .bak, .pv, .db, .DS_Store, temp files (~$...)
"""

import argparse
import email
import logging
import os
import sys
from email import policy
from pathlib import Path

import fitz  # PyMuPDF
import openpyxl
from docx import Document as DocxDocument

try:
    import extract_msg
except ImportError:
    extract_msg = None

logging.basicConfig(
    level=logging.INFO,
    format="%(levelname)-8s %(message)s",
)
log = logging.getLogger(__name__)

SKIP_EXTENSIONS = {".dwg", ".zip", ".bak", ".pv", ".db", ".xls", ".xlsb", ".xltx"}
SKIP_NAMES = {".DS_Store", "Thumbs.db", "convert.py", "requirements.txt"}
SKIP_DIRS = {".venv", "venv", "output", "__pycache__", ".git", "node_modules"}


# ---------------------------------------------------------------------------
# Converters
# ---------------------------------------------------------------------------

def convert_pdf(path: Path) -> str:
    doc = fitz.open(str(path))
    parts: list[str] = [f"# {path.stem}\n"]
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        if text:
            parts.append(f"## Page {i + 1}\n\n{text}")
    doc.close()
    if len(parts) == 1:
        parts.append("*(No extractable text — likely a scanned/image PDF)*")
    return "\n\n".join(parts)


def convert_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts: list[str] = [f"# {path.stem}\n"]

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            para = next(
                (p for p in doc.paragraphs if p._element is element), None
            )
            if para is None:
                continue
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower()
            if "heading 1" in style:
                parts.append(f"## {text}")
            elif "heading 2" in style:
                parts.append(f"### {text}")
            elif "heading 3" in style:
                parts.append(f"#### {text}")
            else:
                parts.append(text)

        elif tag == "tbl":
            table = next(
                (t for t in doc.tables if t._element is element), None
            )
            if table is None:
                continue
            parts.append(_table_to_md(
                [[cell.text.strip() for cell in row.cells] for row in table.rows]
            ))

    return "\n\n".join(parts)


def convert_xlsx(path: Path) -> str:
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = [f"# {path.stem}\n"]

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(cells)
        if rows:
            parts.append(f"## {sheet_name}\n")
            parts.append(_table_to_md(rows))

    wb.close()
    return "\n\n".join(parts)


def convert_eml(path: Path) -> str:
    with open(path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=policy.default)

    parts: list[str] = [f"# {path.stem}\n"]
    parts.append(f"- **From:** {msg.get('From', 'N/A')}")
    parts.append(f"- **To:** {msg.get('To', 'N/A')}")
    parts.append(f"- **Date:** {msg.get('Date', 'N/A')}")
    parts.append(f"- **Subject:** {msg.get('Subject', 'N/A')}")
    parts.append("")

    body = msg.get_body(preferencelist=("plain", "html"))
    if body:
        content = body.get_content()
        parts.append(content.strip() if isinstance(content, str) else content.decode("utf-8", errors="replace").strip())

    return "\n".join(parts)


def convert_msg(path: Path) -> str:
    if extract_msg is None:
        return f"# {path.stem}\n\n*(extract-msg not installed — run `pip install extract-msg`)*"

    msg = extract_msg.Message(str(path))
    parts: list[str] = [f"# {path.stem}\n"]
    parts.append(f"- **From:** {msg.sender or 'N/A'}")
    parts.append(f"- **To:** {msg.to or 'N/A'}")
    parts.append(f"- **Date:** {msg.date or 'N/A'}")
    parts.append(f"- **Subject:** {msg.subject or 'N/A'}")
    parts.append("")
    if msg.body:
        parts.append(msg.body.strip())
    msg.close()
    return "\n".join(parts)


def convert_text(path: Path) -> str:
    text = path.read_text(errors="replace").strip()
    return f"# {path.stem}\n\n{text}"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _table_to_md(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    n_cols = max(len(r) for r in rows)
    # pad rows to same width
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    header = rows[0]
    separator = ["---"] * n_cols
    body = rows[1:]

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


CONVERTERS = {
    ".pdf": convert_pdf,
    ".docx": convert_docx,
    ".xlsx": convert_xlsx,
    ".xlsm": convert_xlsx,
    ".eml": convert_eml,
    ".msg": convert_msg,
    ".txt": convert_text,
    ".log": convert_text,
}


def should_skip(path: Path) -> bool:
    name = path.name
    if name in SKIP_NAMES:
        return True
    if name.startswith("~") or name.startswith("~$"):
        return True
    if path.suffix.lower() in SKIP_EXTENSIONS:
        return True
    return False


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def process_directory(input_dir: Path, output_dir: Path, *, dry_run: bool = False):
    stats = {"converted": 0, "skipped": 0, "failed": 0, "no_converter": 0}

    for root, dirs, files in os.walk(input_dir):
        dirs[:] = [d for d in dirs if d not in SKIP_DIRS]
        root_path = Path(root)
        for filename in sorted(files):
            src = root_path / filename
            rel = src.relative_to(input_dir)

            if should_skip(src):
                log.debug("SKIP   %s", rel)
                stats["skipped"] += 1
                continue

            ext = src.suffix.lower()
            converter = CONVERTERS.get(ext)
            if converter is None:
                log.debug("NOCONV %s (unsupported %s)", rel, ext)
                stats["no_converter"] += 1
                continue

            md_name = f"{src.stem} ({ext.lstrip('.')}){'.md'}"
            dest = output_dir / rel.parent / md_name

            if dry_run:
                log.info("DRYRUN %s → %s", rel, dest.relative_to(output_dir))
                stats["converted"] += 1
                continue

            try:
                markdown = converter(src)
                dest.parent.mkdir(parents=True, exist_ok=True)
                dest.write_text(markdown, encoding="utf-8")
                log.info("OK     %s", rel)
                stats["converted"] += 1
            except Exception:
                log.exception("FAIL   %s", rel)
                stats["failed"] += 1

    log.info("---")
    log.info(
        "Done: %d converted, %d skipped, %d failed, %d unsupported",
        stats["converted"],
        stats["skipped"],
        stats["failed"],
        stats["no_converter"],
    )
    return stats


def main():
    parser = argparse.ArgumentParser(
        description="Convert sizing-automation documents to Markdown"
    )
    parser.add_argument(
        "--input",
        type=Path,
        default=Path(__file__).parent,
        help="Input directory (default: this script's directory)",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path(__file__).parent / "output",
        help="Output directory for Markdown files (default: ./output)",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="List files that would be converted without actually converting",
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Show debug-level logs (skipped files, etc.)",
    )
    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    log.info("Input:  %s", args.input.resolve())
    log.info("Output: %s", args.output.resolve())

    process_directory(args.input, args.output, dry_run=args.dry_run)


if __name__ == "__main__":
    main()
