import os
from pathlib import Path


def extract_docx_text(filepath: str) -> str:
    """Extract all text from a .docx file including tables."""
    from docx import Document

    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


PROJECT_DOCS = {
    "proj-gas-dryer": [
        ("PSR_GasDryer_NES-AGP-2026-007_v2.docx", "Project Specific Requirements - Gas Dehydration Unit"),
        ("RFQ_GasDryer_NES-AGP-2026-007_v2.docx", "RFQ - Gas Dehydration Package Unit"),
    ],
    "proj-h2pdu": [
        ("PSR_H2PDU_MPE-AGHP-2026-003.docx", "Project Specific Requirements - H2 Purification & Drying Unit"),
        ("RFQ_H2PDU_MPE-AGHP-2026-003.docx", "RFQ - H2 Purification & Drying Unit Package"),
    ],
    "proj-hyros": [],
    "proj-asco": [],
}

_doc_cache: dict[str, str] = {}


def load_project_documents(project_id: str, workspace_root: str) -> str:
    """Load and concatenate all documents for a given project."""
    cache_key = project_id
    if cache_key in _doc_cache:
        return _doc_cache[cache_key]

    docs = PROJECT_DOCS.get(project_id, [])
    if not docs:
        return ""

    parts = []
    for filename, display_name in docs:
        filepath = os.path.join(workspace_root, filename)
        if os.path.exists(filepath):
            text = extract_docx_text(filepath)
            parts.append(f"=== {display_name} ({filename}) ===\n{text}")

    result = "\n\n".join(parts)
    _doc_cache[cache_key] = result
    return result


def get_available_projects() -> list[dict]:
    """Return list of projects that have backend document support."""
    return [
        {
            "id": "proj-gas-dryer",
            "name": "Adriatic GDU – ARC Phase II",
            "client": "Arcamind Solutions",
            "system_type": "GAS_DEHYDRATION",
            "status": "new",
            "rfq_number": "ARC-AGP-2026-007",
        },
        {
            "id": "proj-h2pdu",
            "name": "Al Dhafra H2 PDU – Phase I",
            "client": "Arcamind",
            "system_type": "H2_PURIFICATION_DRYING",
            "status": "new",
            "rfq_number": "MPE-AGHP-2026-003",
        },
    ]
